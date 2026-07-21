"""Generate offline PDF and DOCX contract fixtures for manual API testing."""

from pathlib import Path

import fitz
from docx import Document


OUTPUT = Path(__file__).resolve().parent


def create_pdf() -> None:
    pdf = fitz.open()
    pages = [
        "采购合同\n\n合同编号：HT-2026-001\n甲方：北京甲方有限公司",
        "乙方：上海乙方有限公司\n\n合同金额：人民币100,000.00\n"
        "签订日期：2026年7月18日\n\n第一条 双方约定按期履约。",
    ]
    for text in pages:
        page = pdf.new_page()
        page.insert_textbox(
            fitz.Rect(72, 72, 520, 760),
            text,
            fontsize=12,
            fontname="china-s",
        )
    pdf.save(OUTPUT / "sample_contract.pdf")
    pdf.close()


def create_docx() -> None:
    document = Document()
    document.add_heading("采购合同", level=0)
    document.add_heading("合同信息", level=1)
    document.add_paragraph("合同编号：HT-2026-001")
    table = document.add_table(rows=4, cols=2)
    values = [
        ("甲方", "北京甲方有限公司"),
        ("乙方", "上海乙方有限公司"),
        ("合同金额", "人民币100,000.00"),
        ("签订日期", "2026年7月18日"),
    ]
    for row, (key, value) in zip(table.rows, values, strict=True):
        row.cells[0].text = key
        row.cells[1].text = value
    document.add_heading("合同正文", level=1)
    document.add_paragraph("第一条 双方约定按期履约。")
    document.save(OUTPUT / "sample_contract.docx")


if __name__ == "__main__":
    create_pdf()
    create_docx()
    print(f"Created sample files under {OUTPUT}")

