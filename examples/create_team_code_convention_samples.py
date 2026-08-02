"""Generate the native DOCX and image-only PDF used by the OCR acceptance test."""

from __future__ import annotations

import io
from pathlib import Path

import fitz
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
FONT = Path("/System/Library/Fonts/Hiragino Sans GB.ttc")
PAGES = [
    [
        "1. Java 代码规范",
        "类名必须使用大驼峰命名。",
        "单个方法不得超过80行。",
        "禁止使用System.out.println。",
        "不得硬编码密码或Token。",
        "单元测试覆盖率不得低于80%。",
        "表格规则",
        "规则 | 要求 | 示例",
        "类名 | 必须使用大驼峰 | UserService",
        "方法长度 | 不得超过80行 | calculateTotal",
    ],
    [
        "2. 正确示例与错误示例",
        "正确示例：public class UserService {}",
        "错误示例：public class user_service {}",
        "正确示例：logger.info(\"created\");",
        "错误示例：System.out.println(\"created\");",
        "安全要求：不得硬编码密码或Token。",
        "测试要求：单元测试覆盖率不得低于80%。",
    ],
]


def create_docx(path: Path) -> None:
    document = Document()
    document.styles["Normal"].font.name = "Arial Unicode MS"
    document.styles["Normal"].font.size = Pt(11)
    document.add_heading("团队 Java 代码规范", 0)
    document.add_heading("1. Java 代码规范", 1)
    for text in PAGES[0][1:6]:
        document.add_paragraph(text, style="List Bullet")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, ["规则", "要求", "示例"]):
        cell.text = value
    for values in (["类名", "必须使用大驼峰", "UserService"], ["方法长度", "不得超过80行", "calculateTotal"]):
        for cell, value in zip(table.add_row().cells, values):
            cell.text = value
    document.add_page_break()
    document.add_heading("2. 正确示例与错误示例", 1)
    code_style = document.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Menlo"
    code_style.font.size = Pt(10)
    for text in PAGES[1][1:]:
        document.add_paragraph(text, style="Code")
    document.save(path)


def _font(size: int):
    if FONT.is_file():
        return ImageFont.truetype(str(FONT), size)
    return ImageFont.truetype("DejaVuSans.ttf", size)


def _page_image(lines: list[str]) -> Image.Image:
    image = Image.new("RGB", (2480, 3508), "white")
    draw = ImageDraw.Draw(image)
    title_font = _font(58)
    body_font = _font(46)
    y = 170
    for index, text in enumerate(lines):
        font = title_font if index == 0 or text == "表格规则" else body_font
        draw.text((160, y), text, fill="black", font=font)
        y += 130 if index == 0 else 105
    return image


def create_scanned_pdf(path: Path) -> None:
    pdf = fitz.open()
    try:
        for lines in PAGES:
            image = _page_image(lines)
            stream = io.BytesIO()
            image.save(stream, format="PNG")
            page = pdf.new_page(width=595, height=842)
            page.insert_image(page.rect, stream=stream.getvalue())
        pdf.save(path, deflate=True)
    finally:
        pdf.close()


def main() -> None:
    create_docx(ROOT / "team_code_convention.docx")
    create_scanned_pdf(ROOT / "team_code_convention_scanned.pdf")


if __name__ == "__main__":
    main()
