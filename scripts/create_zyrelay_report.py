"""Build the verified ZYRelay PaddleOCR capability and demo report.

The report is generated from the checked-in demonstration outputs rather than
invented benchmark data.  It is intentionally a stakeholder-facing document:
technical enough to explain traceability, concise enough to use in a review.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "ZYRelay_txt模型能力与演示说明.docx"

# standard_business_brief preset — all values are encoded here deliberately.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
GREEN = "1F6B4F"
GOLD = "7A5A00"
RED = "9B1C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float], indent_dxa: int = 120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for i, width in enumerate(widths_in):
        grid.gridCol_lst[i].set(qn("w:w"), str(round(width * 1440)))
    for row in table.rows:
        set_row_cant_split(row)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[idx])
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(round(widths_in[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_row_cant_split(row) -> None:
    """Keep a logical table row together when it reaches a page boundary."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_run(run, size: float = 11, color: str = "000000", bold: bool = False, italic: bool = False, font: str = "Calibri") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(paragraph, text: str, **kwargs):
    run = paragraph.add_run(text)
    set_run(run, **kwargs)
    return run


def set_para(paragraph, before: float = 0, after: float = 6, line: float = 1.10, align=None) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_body(doc, text: str, before: float = 0, after: float = 6, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    set_para(p, before, after)
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, bold=True)
        add_text(p, text[len(bold_prefix):])
    else:
        add_text(p, text)
    return p


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph()
    set_para(p, 0, 8, 1.167)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.25)
    add_text(p, "• ")
    add_text(p, text)


def add_heading(doc, text: str, level: int) -> None:
    p = doc.add_paragraph()
    if level == 1:
        set_para(p, 16, 8, 1.10)
        add_text(p, text, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_para(p, 12, 6, 1.10)
        add_text(p, text, size=13, color=BLUE, bold=True)
    else:
        set_para(p, 8, 4, 1.10)
        add_text(p, text, size=12, color=DARK_BLUE, bold=True)
    set_keep_with_next(p)


def add_callout(doc, title: str, text: str, accent: str = DARK_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    set_para(p, 0, 1, 1.10)
    add_text(p, title + "  ", size=11, color=accent, bold=True)
    add_text(p, text, size=11, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], font_size: float = 9.4) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        set_para(p, 0, 0, 1.0)
        add_text(p, text, size=font_size, color=INK, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for idx, text in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            set_para(p, 0, 0, 1.05)
            add_text(p, text, size=font_size, color="000000")
    p = doc.add_paragraph()
    set_para(p, 4, 4, 1.0)
    add_text(p, "数据来源：本项目离线模型验收与示例输出（2026-08-02）。", size=8.5, color=MUTED, italic=True)


def add_code_block(doc, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    p = cell.paragraphs[0]
    set_para(p, 0, 0, 1.0)
    run = p.add_run(text)
    set_run(run, size=8.8, color="263238", font="Menlo")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(paragraph, "ZYRelay DocIntelligence  |  ", size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    set_para(header, 0, 0, 1.0)
    add_text(header, "ZYRELAY DOCINTELLIGENCE  /  技术能力与演示说明", size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    set_para(footer, 0, 0, 1.0)
    add_page_number(footer)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    set_para(p, 20, 4, 1.0)
    add_text(p, "技术验证说明", size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    set_para(p, 0, 4, 1.0)
    add_text(p, "ZYRelay DocIntelligence", size=24, color=INK, bold=True)
    p = doc.add_paragraph()
    set_para(p, 0, 14, 1.0)
    add_text(p, "PaddleOCR 模型能力、演示与结果", size=15, color=DARK_BLUE)
    for label, value in [
        ("面向对象", "管理汇报与技术评审"),
        ("版本", "1.0.0（企业级可追溯语义对象层）"),
        ("日期", "2026 年 8 月 2 日"),
        ("验证范围", "PDF / DOCX 解析、扫描 PDF OCR、代码规范结构化与可追溯输出"),
    ]:
        p = doc.add_paragraph()
        set_para(p, 0, 2, 1.0)
        add_text(p, label + "：", size=10.5, color=INK, bold=True)
        add_text(p, value, size=10.5, color="000000")
    p = doc.add_paragraph()
    set_para(p, 10, 6, 1.0)
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), BLUE)
    pbdr.append(bottom)
    p_pr.append(pbdr)


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    add_title_block(doc)

    add_heading(doc, "1. 结论摘要", 1)
    add_body(doc, "ZYRelay DocIntelligence 已在本地 CPU 环境完成真实 PaddleOCR 集成与离线验收。系统可将 PDF、DOCX 中的非结构化文本，特别是扫描 PDF 的图像文字，转换为可查询、可定位、可追溯的标准化 UOM Package。")
    add_callout(doc, "本次核心成果", "扫描版 2 页 PDF 已真实调用 PaddleOCR 3.7.0，识别出 17 条 OCR 文本行、提取 9 条团队代码规范候选，并为每条候选保留页码、文本块、字符偏移、坐标框、置信度、模型执行记录和 provenance 链路。", GREEN)
    add_bullet(doc, "规则优先：标签、规则表达式和业务对象候选均有原文证据；LLM 不作为主流程依赖。")
    add_bullet(doc, "离线可控：模型由管理员预置并校验，正常文档处理禁止自动下载模型。")
    add_bullet(doc, "可扩展：同一能力通过 Python、HTTP、CLI 三种入口调用，OCR 以资源插件方式接入。")

    add_heading(doc, "2. 模型与系统能力", 1)
    add_table(doc,
        ["能力", "实现方式", "结构化输出 / 管理价值"],
        [
            ["PDF / DOCX 解析", "PyMuPDF、python-docx；按页、段落、标题、表格生成统一 block。", "获得有序文本块；DOCX、文本型 PDF 不必启用 OCR。"],
            ["扫描 PDF 识别", "OCR Gate 先判断原生文本；仅对无可提取文本且含图像的页面，用 PyMuPDF 渲染后调用 PaddleOCR。", "生成 OCR block，携带页码、bbox、polygon、阅读顺序与 OCR 置信度。"],
            ["标签与语义索引", "YAML 驱动标签；正则、别名精确匹配和可选模糊匹配。", "按 label_code 建 semantic_index，可定位至 block、页码和 offset。"],
            ["代码规范提取", "配置与规则驱动的候选生成器，识别命名、日志、安全、测试等规范。", "输出可执行或可审核的 rule_expression，不将推断当作事实。"],
            ["证据与追溯", "Provenance 串联候选、mention、block、OCR 坐标、模型执行、资源计划与 Ground Snapshot。", "可回到“第几页、哪一行/区域、由哪个模型处理”进行人工复核。"],
            ["标准交付", "MOM / SOM / BOM 统一 UOM JSON Package，本地原子写入。", "便于后续接入检索、审核、知识治理或质量平台。"],
        ], [1.35, 2.55, 2.60], 9.0)

    add_heading(doc, "3. 处理架构与边界", 1)
    add_body(doc, "系统遵循“先可解释、后智能化”的处理路径。OCR 是可替换资源插件，不和核心解析逻辑绑定；LLM 增强默认关闭，且不得覆盖高置信度的规则结果。")
    add_code_block(doc, "上传文档  →  文件校验 / GroundChoose  →  资源计划  →  原生解析\n"
                        "                                      ↓\n"
                        "                           OCR Gate（仅扫描 PDF）\n"
                        "                                      ↓\n"
                        "页面渲染（短期 PNG） → PaddleOCR → blocks → 标签 / semantic_index\n"
                        "                                      ↓\n"
                        "代码规范候选 → 证据校验 → provenance → UOM Package / Python / HTTP / CLI")
    add_table(doc,
        ["设计约束", "当前处理方式"],
        [
            ["不伪造识别结果", "模型或缓存不可用时使用 NoOp OCR，明确标记部分完成，不生成虚假文字。"],
            ["不自动联网", "运行参数 allow_download=false、offline_mode=true；预置完成后运行时不下载模型。"],
            ["不做超范围推理", "只生成有证据的 detected 候选；不写入知识图谱，不自动确认业务事实。"],
            ["保留原文定位", "normalized_text 用于匹配；原始 text 和 offset 用于证据，避免标准化破坏定位。"],
        ], [1.60, 4.90], 9.2)

    add_heading(doc, "4. 已验证模型配置", 1)
    add_table(doc,
        ["项目", "验证结果"],
        [
            ["运行环境", "macOS 26.2，Apple Silicon（ARM64），Python 3.13.7，CPU 模式。"],
            ["OCR 运行时", "PaddlePaddle 3.3.1；PaddleOCR 3.7.0。"],
            ["模型版本", "PaddleOCR 3.7.0 / PP-OCRv6-medium。"],
            ["离线模型", "PP-LCNet_x1_0_doc_ori、PP-LCNet_x1_0_textline_ori、PP-OCRv6_medium_det、PP-OCRv6_medium_rec。"],
            ["模型缓存", "data/model_cache/paddleocr；约 146 MB；ready.json 和 12 个 SHA-256 校验条目已生成。"],
            ["图像渲染", "仅对待 OCR 页面以 200 DPI、RGB 生成临时 PNG；输出中只保留 relay:// 工件引用与摘要。"],
            ["离线验证", "models_ready=true、offline_ready=true、verified=true；本地烟雾推理识别到 “PaddleOCR80 System.out println”。"],
        ], [1.65, 4.85], 9.1)

    add_heading(doc, "5. 演示设计", 1)
    add_body(doc, "演示使用同一份《Java 团队代码规范》内容生成两个真实输入样例，以对比原生文本通道和扫描件 OCR 通道。样例覆盖：类名大驼峰、方法不超过 80 行、禁止 System.out.println、禁止硬编码密码或 Token、单元测试覆盖率不低于 80%，并包含表格与正反示例。")
    add_table(doc,
        ["样例", "输入特性", "应走通道", "验收目标"],
        [
            ["team_code_convention.docx", "DOCX，包含段落与表格。", "python-docx 原生解析；OCR 跳过。", "验证段落、表格、标签、代码规范候选和 UOM 输出。"],
            ["team_code_convention_scanned.pdf", "2 页图像型 PDF；PyMuPDF 原生提取文本为空。", "OCR Gate → 200 DPI 页图 → PaddleOCR。", "验证真实 OCR、坐标、置信度、模型执行与 provenance。"],
        ], [1.65, 1.65, 1.70, 1.50], 8.9)

    add_heading(doc, "6. 演示调用方式", 1)
    add_body(doc, "模型安装和文档处理均可在命令行完成；服务化部署后也可通过 HTTP 调用。以下命令为项目内实际支持的标准入口。")
    add_code_block(doc, "# 1) 管理员一次性预置与校验模型（联网仅发生在该阶段）\n"
                        "python -m zyrelay.models install paddleocr\n"
                        "python -m zyrelay.models verify paddleocr\n\n"
                        "# 2) CLI 处理扫描 PDF\n"
                        "zyrelay relay process examples/team_code_convention_scanned.pdf\n\n"
                        "# 3) HTTP：上传后取得 execution_id，再查询执行与 provenance\n"
                        "POST /api/v1/relay/process\n"
                        "GET  /api/v1/relay/executions/{execution_id}\n"
                        "GET  /api/v1/relay/provenance/{provenance_id}")
    add_body(doc, "Python SDK、HTTP 和 CLI 共享同一套 Relay 执行与 UOM 输出契约；接口不会因调用方式不同而改变证据字段。", bold_prefix="Python SDK、HTTP 和 CLI")

    add_heading(doc, "7. 实际演示结果", 1)
    add_callout(doc, "结论", "两种输入均完成处理。扫描 PDF 不是模拟结果：其执行记录显示 PaddleOCR 已真实运行，未使用 fallback，且所有 OCR 文本均可关联到模型执行编号。", GREEN)
    add_table(doc,
        ["指标", "DOCX 原生文本样例", "扫描 PDF OCR 样例"],
        [
            ["执行状态", "completed", "completed"],
            ["执行编号", "EXEC-1C9E838E214B4F9F", "EXEC-7245CF4A6F1748D7"],
            ["页数 / OCR", "DOCX 无可靠物理分页；OCR 跳过。", "2 页；原生文本为空；17 条 OCR 行。"],
            ["代码规范候选", "9 条", "9 条"],
            ["候选置信度", "规则候选 0.94", "规则候选 0.94；OCR 平均置信度 0.9972。"],
            ["OCR 模型执行", "无（按设计跳过）", "MEXEC-2FF28CA59F5B4828；fallback_used=false。"],
            ["处理告警 / 错误", "3 条可解释匹配告警；0 错误。", "3 条可解释告警；0 错误。"],
            ["总体耗时", "221.71 ms", "41,951.60 ms（含模型加载与 2 页 OCR）。"],
        ], [1.45, 2.25, 2.80], 8.8)

    add_heading(doc, "8. 识别结果示例", 1)
    add_table(doc,
        ["原文规则", "结构化规则表达式", "页码 / 证据"],
        [
            ["类名必须使用大驼峰命名。", "target=class_name；matches_regex；PascalCase", "第 1 页，BLK-000002；OCR 置信度 0.999756。"],
            ["单个方法不得超过80行。", "target=function_length；less_than_or_equal；80 lines", "第 1 页，BLK-000003。"],
            ["禁止使用System.out.println。", "target=source_code；not_contains；System.out.println", "第 1 页，BLK-000004。"],
            ["不得硬编码密码或Token。", "target=source_code；not_contains_sensitive_secret", "第 1 页，BLK-000005。"],
            ["单元测试覆盖率不得低于80%。", "target=unit_test_coverage；greater_than_or_equal；80 percent", "第 1 页，BLK-000006。"],
        ], [1.80, 2.60, 2.10], 8.7)

    add_heading(doc, "9. 可追溯性演示", 1)
    add_body(doc, "以“类名必须使用大驼峰命名。”为例，系统不是只返回一个关键词，而是保留完整的证据链。前端或审核人员可依据 provenance_id 查询并回到原始页面区域。")
    add_code_block(doc, "CodeConvention candidate\n"
                        "  CONV-8d122c1fa3bf6b56  (confidence=0.94, status=detected)\n"
                        "      ↓ provenance_id\n"
                        "  PROV-053F5C15985844BE\n"
                        "      ↓ source evidence\n"
                        "  Page 1 / BLK-000002 / 原文：类名必须使用大驼峰命名。\n"
                        "  bbox=[105, 195, 456, 233] / OCR confidence=0.999756\n"
                        "      ↓ model execution\n"
                        "  MEXEC-2FF28CA59F5B4828 / PaddleOCR 3.7.0 / CPU\n"
                        "      ↓ governance context\n"
                        "  Resource Plan RPLAN-9757864EFEBE4702 / Ground Snapshot GSNAP-2A1241015705496D")
    add_body(doc, "这条链路同时满足“关键词可回溯到原文”和“模型处理过程可审计”两项要求。对于 OCR 文本，坐标框可直接用于在页面预览中高亮定位。")

    add_heading(doc, "10. 扫描 PDF 性能记录", 1)
    add_table(doc,
        ["指标", "实测结果"],
        [
            ["模型首次加载", "4,264.92 ms"],
            ["第 1 页 OCR", "19,983.76 ms；10 行；平均置信度 0.998208"],
            ["第 2 页 OCR", "17,090.41 ms；7 行；平均置信度 0.995641"],
            ["OCR 执行总时长", "41,340.35 ms"],
            ["模型置信度汇总", "min 0.980199；max 0.999946；average 0.997151"],
            ["端到端总时长", "41,951.60 ms（含解析、OCR、结构化、索引、provenance 与保存）。"],
        ], [2.10, 4.40], 9.1)
    add_body(doc, "说明：以上为单次本地 CPU 验证结果，用于功能验收和容量估算，不应被视为生产环境 SLA。实际耗时会受页图复杂度、CPU、并发和模型冷启动影响。", bold_prefix="说明：")

    add_heading(doc, "11. 测试与验证", 1)
    add_table(doc,
        ["验证项", "结果", "说明"],
        [
            ["默认单元 / API 测试", "50 passed, 1 skipped", "默认轻量环境不安装 OCR 大运行时；核心功能和接口已覆盖。"],
            ["真实模型集成测试", "1 passed", "在 .venv-paddleocr 中真实调用本地 PaddleOCR；不访问网络。"],
            ["编译检查", "passed", "Python compileall 通过。"],
            ["代码覆盖率", "86%", "默认测试环境统计；新增重模型路径由独立 integration 标记验证。"],
            ["模型预置验证", "verified=true", "四个官方模型就绪、烟雾推理成功、校验记录存在。"],
        ], [1.85, 1.25, 3.40], 9.0)

    add_heading(doc, "12. 输出物与对接方式", 1)
    add_table(doc,
        ["交付物", "位置 / 接口", "用途"],
        [
            ["完整 Relay 结果", "examples/output/scanned_pdf_relay_result.json", "包含 document、blocks、labels、semantic_index、conventions、business_objects、UOM。"],
            ["模型执行记录", "examples/output/scanned_pdf_model_execution.json", "审计模型版本、耗时、每页统计、工件引用和 fallback 状态。"],
            ["Provenance 样例", "examples/output/scanned_pdf_provenance.json", "查询“候选 → 原文 → OCR → 模型 → 配置”的完整链路。"],
            ["服务接口", "/api/v1/relay/process、/executions/{id}、/provenance/{id}", "供上层平台上传、查询和回溯。"],
        ], [1.60, 2.50, 2.40], 8.8)

    add_heading(doc, "13. 当前限制与下一步", 1)
    add_table(doc,
        ["当前限制", "建议下一步"],
        [
            ["OCR 已完成 CPU 离线 MVP 验证，尚未形成批量并发性能基线。", "增加样本文档集与并发压测，明确冷启动、热启动和页数分档指标。"],
            ["DOCX 缺少可靠的物理页码，当前保留逻辑 block 与 offset。", "如需 Word 页码级定位，可在展示层引入固定版式渲染或页码映射策略。"],
            ["代码规范提取以配置、词典和规则为主，面向明确文本。", "按语言和团队扩充 labels / Ground Truth；规则保持版本化与人工审核。"],
            ["敏感字段规则可被识别，但未连接组织级合规策略。", "接入审核工作流，增加 pending / accepted / rejected 人工确认闭环。"],
            ["LLM 仅预留可选增强接口，默认关闭。", "在真实 Ground Truth 基础上逐步启用，并强制引用 block_id 与原文证据。"],
        ], [2.75, 3.75], 8.9)

    add_heading(doc, "附录：汇报口径", 1)
    add_callout(doc, "一句话表述", "ZYRelay 已把团队代码规范文档从“可阅读的非结构化内容”转为“可检索、可执行、可回到原文核验”的结构化资产；扫描 PDF 场景已通过真实本地 PaddleOCR 完成离线验证。", DARK_BLUE)
    add_body(doc, "建议汇报时强调三点：一是模型真实运行且不依赖云端；二是识别结果保留页码、坐标和模型审计信息；三是当前为可控 MVP，后续可在不推翻架构的前提下扩展 OCR、规则库和审核流程。")

    # The core properties make the artifact usable in document management systems.
    props = doc.core_properties
    props.title = "ZYRelay DocIntelligence — PaddleOCR 模型能力与演示说明"
    props.subject = "真实离线 OCR、代码规范结构化与可追溯演示"
    props.author = "ZYRelay 项目组"
    props.keywords = "ZYRelay, PaddleOCR, 文档智能, OCR, 代码规范, 可追溯"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
