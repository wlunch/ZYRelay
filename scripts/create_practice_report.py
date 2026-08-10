"""Create the ZYRelay practice report DOCX with a deterministic Word layout."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "ZYRelay_DocIntelligence_实践报告.docx"
CHART = ROOT / "docs" / "zyrelay_practice_metrics.png"


def set_font(run, size: float, *, bold: bool = False, east_asia: str = "Songti SC"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph(paragraph, *, before=0, after=0, first_line=0, align=None):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Cm(first_line) if first_line else None
    if align is not None:
        paragraph.alignment = align


def add_field(paragraph, field: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, 10.5)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, title in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Cm(widths[i])
        set_cell_shading(cell, "D9EAF7")
        p = cell.paragraphs[0]
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(title)
        set_font(r, 10.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Cm(widths[i])
            p = cells[i].paragraphs[0]
            set_paragraph(p)
            r = p.add_run(str(value))
            set_font(r, 10.5)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph(p, before=10 if level == 1 else 6, after=4)
    r = p.add_run(text)
    set_font(r, 15 if level == 1 else 13, bold=True, east_asia="STHeiti")
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, after=2, first_line=0.74, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    r = p.add_run(text)
    set_font(r, 11.5)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, before=2, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(text)
    set_font(r, 10.5)


def make_chart():
    try:
        import matplotlib.pyplot as plt
    except ImportError:
        from PIL import Image, ImageDraw, ImageFont

        image = Image.new("RGB", (1260, 600), "white")
        draw = ImageDraw.Draw(image)
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 30)
        title = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 34)
        labels = ["Offline tests", "Evidence", "Provenance", "Coverage"]
        values = [100, 100, 100, 84]
        colors = ["#4F81BD", "#5B9BD5", "#70AD47", "#ED7D31"]
        draw.text((400, 30), "ZYRelay v1.0 Validation Metrics", fill="black", font=title)
        draw.line((120, 500, 1160, 500), fill="#333333", width=3)
        for index, (label, value, color) in enumerate(zip(labels, values, colors, strict=True)):
            x = 170 + index * 250
            height = int(value * 3.5)
            draw.rectangle((x, 500 - height, x + 130, 500), fill=color)
            draw.text((x + 35, 460 - height), f"{value}%", fill="black", font=font)
            draw.text((x - 20, 530), label, fill="black", font=font)
        image.save(CHART)
        return True
    labels = ["离线测试通过", "证据校验", "溯源校验", "代码覆盖率"]
    values = [100, 100, 100, 84]
    fig, ax = plt.subplots(figsize=(7.0, 3.4))
    bars = ax.bar(labels, values, color=["#4F81BD", "#5B9BD5", "#70AD47", "#ED7D31"])
    ax.set_ylim(0, 110)
    ax.set_ylabel("百分比（%）")
    ax.set_title("ZYRelay v1.0 实践验证指标")
    for bar, value in zip(bars, values, strict=True):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 2, f"{value}%", ha="center", fontsize=10)
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(CHART, dpi=180)
    plt.close(fig)
    return True


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    normal.font.size = Pt(11.5)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)
    for name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        style.font.size = Pt(size)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(fp, "PAGE")

    # Cover
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("实践报告")
    set_font(r, 28, bold=True, east_asia="STHeiti")
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("ZYRelay DocIntelligence 文档智能平台\n基于 0.6 版本的迭代优化实践")
    set_font(r, 22, bold=True, east_asia="STHeiti")
    for _ in range(8):
        doc.add_paragraph()
    for label, value in [
        ("课题名称", "ZYRelay DocIntelligence 文档智能平台实践"),
        ("实践单位", "____________________________"),
        ("团队成员", "____________________________（姓名+学号+院系）"),
        ("实践指导教师", "____________________________"),
        ("完成日期", "2026 年 8 月"),
    ]:
        p = doc.add_paragraph()
        set_paragraph(p, after=7, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(f"{label}：{value}")
        set_font(r, 14)
    doc.add_page_break()

    # Abstract
    add_heading(doc, "摘  要", 1)
    abstract = (
        "本实践围绕企业 PDF、DOCX 文档的结构化处理需求，持续建设 ZYRelay DocIntelligence。"
        "系统的基本职责是将合同、制度、团队代码规范等非结构化文本转换为可验证、可追溯的标准 JSON Package，"
        "为上层检索、代码评审、知识服务和业务系统提供稳定输入，而不承担智能体决策、知识图谱存储或复杂推理。"
        "在 0.6 版本已具备本地 OCR、版面分析、表格识别、语言识别、实体识别和代码块检测等资源插件的基础上，"
        "本次实践进一步完善模型路由、企业资源配置、语义对象、对象级 provenance、插件治理、配置版本控制、"
        "运行记录、基准测试和部署能力。系统坚持规则优先：标签词典、正则规则和 Ground 配置是权威来源；"
        "模型只提供辅助元数据，模型不可用时以启发式资源回退，并将原因写入执行记录。"
        "测试表明，当前离线测试共 75 项通过；合同基准的 6 个案例均完成或以受控降级完成；证据校验和 provenance 校验均为 100%。"
        "同时，项目如实保留了覆盖率 84%、扫描件场景的语义完整率和预期命中率仍需提高等问题。"
        "实践结果说明，面向企业文档智能的首要工作不是追求不可解释的生成能力，而是先建立可配置、可定位、可复核的处理链路。"
    )
    add_body(doc, abstract)
    p = doc.add_paragraph()
    set_paragraph(p, before=4, after=6)
    r = p.add_run("关键词：")
    set_font(r, 11.5, bold=True)
    r = p.add_run("文档智能；规则优先；语义对象；资源插件；原文追溯；UOM Package")
    set_font(r, 11.5)

    add_heading(doc, "1  课题背景与实践目标")
    for text in [
        "企业日常积累的合同、采购文件、制度、操作手册和团队开发规范，大多以 PDF、DOCX 等文本型非结构化文件保存。"
        "传统文件管理通常只能按文件名或全文关键词检索，无法稳定回答某个合同编号出现在哪一页、某项代码规则依据是什么、"
        "某个实体或金额是否来自可复核的原文。对于后续的审批、审计、研发治理和知识服务而言，缺少原文证据的结构化结果价值有限。",
        "ZYRelay 的定位是企业文档智能的基础层。它不直接替代业务系统，也不将文档内容自动写入知识图谱。"
        "它负责把文件解析为页面和逻辑块，按 YAML 配置识别标签，形成语义索引和业务对象候选，并在每个结果上保留页码、"
        "文本块、起止 offset、规则、Ground 快照、资源计划和执行记录。这样，上层系统可以消费结构化对象，"
        "同时仍能回到原始文本进行人工核对。",
        "本实践的目标包括四个方面：第一，保持 PDF、DOCX 的稳定解析与统一 block 输出；第二，以规则和词典为主完成字段、"
        "实体和代码规范提取；第三，将模型能力限制在 OCR、语言、版面、表格、NER 等辅助环节；第四，"
        "建立从配置选择、资源调用到 UOM 输出的完整记录，使处理结果可复现、可解释、可维护。",
    ]:
        add_body(doc, text)

    add_heading(doc, "2  0.6 版本基础与本次实践范围")
    for text in [
        "0.6 版本已经完成了本地轻量资源插件的接入。系统为 OCR、文档分类、语言检测、版面分析、表格识别、"
        "代码块检测、拼写校正和 NER 提供统一的资源接口。每个资源可以报告可用性、健康状态、版本和元数据，"
        "并在执行时生成 ModelExecution、ResourcePlan 和 Provenance。对于无法安装或无法加载的模型，"
        "系统使用 NoOp 或启发式实现，保证文档解析和规则提取不因辅助模型异常而整体失败。",
        "0.6 的重点是“模型资源可以接入流水线”。但是，单纯接入模型仍不足以支持真实团队使用。"
        "如果没有清晰的选择条件，同一份 DOCX 可能无意义地触发 OCR；如果没有资源范围和版本记录，"
        "不同企业、团队或环境得到的结果难以说明差异；如果没有稳定语义对象和对象级证据，"
        "下游系统只能处理松散的 JSON 字段，无法建立可靠的数据消费边界。",
        "因此，本次在 0.6 之后的优化没有推翻原有流水线，而是在固定架构内补充治理能力。"
        "主路径仍然是 GroundChoose、ResourcePlan、解析器、AI 资源插件、规则引擎、语义对象、业务对象、"
        "证据验证、Provenance 和 UOM Package。新增内容的原则是字段可选、配置可扩展、结果可追溯，"
        "不改变既有上传、查询和插件调用接口的基本行为。",
    ]:
        add_body(doc, text)

    add_heading(doc, "3  总体架构与处理流程")
    add_body(doc, "系统采用固定同步流水线，而不是工作流引擎。GroundChoose 先根据企业、团队、项目、模式和文档类型选择规则配置；"
             "ResourcePlan 再根据配置选择解析器、OCR 和辅助模型；解析器生成原始页面和元素；模型只在满足条件时执行；"
             "规则引擎在统一 blocks 上完成标签识别；最后生成 semantic_index、业务对象候选、语义对象、provenance 和 UOM。")
    add_table(doc, ["阶段", "主要输入", "主要输出", "控制原则"], [
        ["GroundChoose", "企业、团队、项目、模式", "Ground 快照", "配置优先、版本冻结"],
        ["ResourcePlan", "资源 YAML、健康状态", "资源绑定与 fallback", "可用性和兼容性检查"],
        ["Parser / OCR", "PDF、DOCX、扫描页", "页面、元素、原文块", "保留原始文本和坐标"],
        ["Rule Engine", "blocks、标签、规则", "mentions、semantic_index", "正则和词典优先"],
        ["Semantic / UOM", "候选、证据、执行记录", "语义对象和标准包", "所有对象可回溯"],
    ], [3.0, 4.2, 4.2, 5.0])
    add_caption(doc, "表 1  ZYRelay 固定处理流程及其输出")
    add_body(doc, "在文本标准化环节，系统采用 Unicode NFKC、换行和连续空格统一等可解释处理，同时保留 block.text 作为证据原文。"
             "匹配使用的 normalized_text 与 evidence 使用的原始 text 分离，offset 始终以原始文本为准。"
             "这样即使标准化改变了空格或字符形式，也不会破坏“关键词位于第几页、第几个文本块、第几个字符”的追溯能力。")

    add_heading(doc, "4  基于 0.6 的具体优化内容")
    add_heading(doc, "4.1 模型路由与执行门控", 2)
    for text in [
        "后续版本增加了 ModelRouter。它不是新的调度系统，而是固定流水线中的判断器：在每个可选模型执行前，"
        "根据文档类型、是否为扫描件、是否已有规则结果、请求开关和资源健康状态，返回“执行”或“跳过”的决定。"
        "例如 DOCX 默认不触发 OCR 和视觉版面模型；原生文本 PDF 不触发 OCR；只有扫描型 PDF 且 enable_ocr 为真时才进入 OCR；"
        "拼写校正仅处理 OCR 文本；已由规则标签稳定覆盖的实体可以不调用 NER。",
        "路由结果不是内部临时判断，而是写入 ModelExecution 和 ResourcePlan。对于每个能力，记录 planned_execution、"
        "actual_execution、skip_reason、gate_decision、输入信号、资源版本和延迟。这样，用户可以区分“模型没有安装”、"
        "“模型被禁用”和“模型按规则不需要执行”，避免把所有缺少模型输出的问题都误判为系统异常。",
    ]:
        add_body(doc, text)

    add_heading(doc, "4.2 企业资源管理与 YAML 覆盖", 2)
    for text in [
        "资源配置从单一默认文件扩展为企业范围配置。ResourcePlan 现在包含 enterprise_id、department_id、team_id、"
        "project_id 和 environment 字段，环境支持 dev、test、prod。资源选择采用确定的覆盖顺序：默认配置、企业配置、"
        "环境覆盖、部门覆盖、团队覆盖、项目覆盖。后一个范围只覆盖明确声明的能力，不会要求用户复制完整默认配置。",
        "以 OCR 为例，生产环境可以配置 PaddleOCR 为主资源、NoOpOCR 为回退；测试环境可以将视觉版面或表格能力切换为"
        "启发式资源，以缩短测试时间和减少模型依赖。每次处理都保留实际选中的资源、主备列表、资源健康状态、"
        "配置版本和配置 hash。因此，当两个环境输出不同结果时，可以先比较 ResourcePlan，而不是只比较最终 JSON。",
    ]:
        add_body(doc, text)

    add_heading(doc, "4.3 稳定语义对象与业务对象候选", 2)
    for text in [
        "原有输出包含 labels、mentions、semantic_index 和 candidates，已经能完成合同编号、金额、日期和主体等定位。"
        "在此基础上，系统新增 semantic_objects 区块，将已有信息组织为 document_object、evidence、entity、"
        "observation、rule、relation、event 和 business_object 等统一对象。这里的“对象”不是知识图谱节点，"
        "而是可独立交换的、证据优先的数据记录。",
        "对象 ID 由文档稳定信息、对象类型、归一化名称、页码、block 和 offset 等来源信息计算。相同文件重复处理时，"
        "同一原文位置的对象保持相同 ID；语义迁移只补充 schema_version，不重新计算 ID。"
        "对于关系和业务对象，系统只在存在明确 evidence 或 mention 来源时生成，不推断文档中没有出现的字段。",
    ]:
        add_body(doc, text)

    add_heading(doc, "4.4 完整 provenance 与原文追溯", 2)
    for text in [
        "对象级 provenance 是本次实践的重要补充。每个语义对象可关联 document_id、page、block_id、"
        "offset、evidence_ids、Ground snapshot、ResourcePlan、模型执行 ID、规则 ID 和创建时间。"
        "例如合同编号的 observation 可以返回匹配文本、页码、block、字符起止位置和 regex 匹配方式；"
        "代码规范 rule 则可以返回其原始段落、规则表达式、适用语言、建议工具和证据验证记录。",
        "这种设计使 semantic_index 不再只是“关键词到文档”的索引，而是“标签代码到可定位证据”的索引。"
        "查询接口支持按 label_code、document_id 和 value 检索，并返回原文证据。上层系统可直接引用 provenance_id，"
        "在出现争议时回到原文，而不必相信无法解释的模型结论。",
    ]:
        add_body(doc, text)

    add_heading(doc, "4.5 插件 SDK 与运行治理", 2)
    for text in [
        "插件层继续保持 Python、HTTP 和 CLI 三种调用方式。插件 manifest 增加了 version、dependencies、"
        "configuration_schema、supported_content_types、supported_languages、license、author 和 permissions 等字段。"
        "服务端新增 health、validate、install、update、disable、enable 等生命周期接口。"
        "其中 install 和 update 只针对已部署、可信的进程内插件进行校验和重新注册，不从网络下载未知可执行代码。",
        "运行记录方面，每个 Relay 执行保存步骤开始结束时间、状态、资源、warning、error、执行历史、重试次数和性能统计。"
        "LLM 保持默认关闭；即使将来接入兼容模型，也只能作为低置信度补充，不能覆盖高置信度规则结果。"
        "对于外部接口，系统避免返回本地内部路径和堆栈信息，上传文件名会规范化处理，文件大小通过配置限制。",
    ]:
        add_body(doc, text)

    add_heading(doc, "5  关键数据模型与接口说明")
    add_body(doc, "UOM Package 顶层保持 schema_version=1.0，包含 source、mom、som、semantic_objects、bom 和 processing。"
             "MOM 保存文档和 blocks；SOM 保存标签、mentions、semantic_index、候选和代码规范；BOM 保存业务对象候选；"
             "processing 保存配置清单、模型路由、性能数据、步骤记录和 warnings。该组织方式使旧调用方继续读取原有字段，"
             "新调用方再按需要使用语义对象和运行信息，避免一次升级破坏已有集成。")
    add_table(doc, ["接口", "用途", "主要返回"], [
        ["POST /api/v1/documents", "上传 PDF 或 DOCX", "document_id、task_id、状态"],
        ["GET /documents/{id}/blocks", "查看逻辑块", "页码、文本、序号、offset"],
        ["GET /documents/{id}/semantic-index", "查看标签语义索引", "label_code、匹配位置和证据"],
        ["GET /documents/{id}/semantic-objects", "查看稳定语义对象", "对象、evidence、provenance"],
        ["POST /api/v1/relay/process", "企业 scope 处理入口", "执行状态、资源计划、UOM 摘要"],
        ["GET /api/v1/plugins/{id}/health", "查看插件状态", "启用状态与 manifest 校验"],
    ], [4.6, 5.2, 6.6])
    add_caption(doc, "表 2  主要接口及返回内容")
    add_body(doc, "在直接展示时，可通过 FastAPI Swagger 页面完成上传、Relay 处理和查询。实际演示宜选择一份含合同编号、"
             "甲乙方、金额、日期和多页正文的合同样例：先展示上传后的 pages 和 blocks，再展示 contract_no 的 semantic_index，"
             "然后打开 semantic_object 和 provenance，最后展示 Contract business_object candidate 的 attributes 来源。"
             "这样可以完整说明系统从原文到结构化结果的过程。")

    add_heading(doc, "6  测试设计与实践结果")
    for text in [
        "测试采用离线方式，不依赖外部模型服务或网络。单元测试覆盖 PDF 文本提取、DOCX 段落与表格提取、"
        "block 顺序、标签正则匹配、别名匹配、offset、semantic_index 聚合、业务对象候选、UOM 序列化、"
        "上传 API、查询 API、插件契约、模型路由、企业资源 scope 和语义对象迁移。"
        "对于 OCR、GLiNER、DocLayout 等可选依赖，测试重点验证资源健康、回退行为和输出格式，而非要求所有平台都下载模型权重。",
        "本次实际运行 .venv-paddleocr 环境中的 pytest -m 'not slow'，离线测试 75 项通过。"
        "slow 标记的真实 OCR 集成用例未纳入这一轮完整回归，因此报告不将其表述为已验证。"
        "Ruff lint 和 format 检查通过；公共数据模型、资源模型、Relay 请求模型和插件生命周期接口的类型检查通过。"
        "完整严格 mypy 仍存在历史类型债务，因此当前 CI 对公共契约实施类型检查，并将全库严格类型治理列为后续工作。",
        "基准测试选取合同类案例执行，6 个案例均完成或以受控 fallback 完成。平均端到端耗时约 4417 毫秒，"
        "范围为 3582 至 5190 毫秒；平均 traced Python 峰值内存约 9.3 MB，范围为 1.1 至 27.4 MB。"
        "证据有效率和 provenance 有效率均为 100%。这些数据仅反映当前本地环境和样例集，不能直接推广为所有文档或所有硬件环境的性能结论。",
    ]:
        add_body(doc, text)
    if make_chart():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(CHART), width=Inches(6.1))
        add_caption(doc, "图 1  当前离线验证中的主要质量指标")
    add_table(doc, ["验证项", "实际结果", "说明"], [
        ["离线回归测试", "75 项通过", "不包含 slow 真实 OCR 集成用例"],
        ["合同基准", "6/6 完成或受控降级", "保留 fallback 和 warning"],
        ["证据有效率", "100%", "结构化结果可回到原文"],
        ["Provenance 有效率", "100%", "对象存在溯源记录"],
        ["全包代码覆盖率", "84%", "未达到 95% 目标"],
        ["预期命中率均值", "75%", "扫描件与 fallback 仍需优化"],
        ["语义证据完整率均值", "67%", "需补充边界场景测试与规则"],
    ], [4.0, 4.0, 8.4])
    add_caption(doc, "表 3  实践验证结果与限制")

    add_heading(doc, "7  结果分析与实践体会")
    for text in [
        "从实践结果看，规则优先方法适合作为企业文档智能的第一层能力。合同编号、金额、日期、主体和规范性表达等内容，"
        "通常具有较稳定的标题、别名或表达形式。通过 YAML 标签、正则和别名识别，可以得到较高置信度和清晰证据。"
        "相比直接让模型给出摘要或结论，这种方式更便于维护：业务人员可以修改配置，研发人员可以测试规则，审计人员可以复核原文。",
        "轻量模型的价值主要在规则难以覆盖的辅助信息。例如 OCR 为扫描页恢复文本，语言检测帮助标注中英文混合块，"
        "代码检测帮助区分配置片段与普通段落，NER 为规则未覆盖的明确实体提供候选。模型结果并不直接成为 Ground Truth，"
        "而是作为 block 元数据或低置信度 observation 保存。这样即使模型缺失、版本不兼容或推理失败，"
        "系统仍能依赖解析器和规则完成基本处理。",
        "实践也发现，最难的问题不是增加模型数量，而是确定模型何时应执行以及结果如何解释。"
        "如果不做门控，模型调用会增加耗时和不稳定性；如果不记录 fallback，用户会误以为模型已经正确执行；"
        "如果不固定对象 ID 和配置 hash，结果难以做回归比较。因此本次把 ModelRouter、ResourcePlan、"
        "Ground Snapshot 和 Provenance 放在同等重要的位置。",
        "目前的指标也说明系统仍处于持续完善阶段。84% 覆盖率低于 95% 目标；扫描件和回退场景会拉低预期命中率与"
        "语义证据完整率。后续优化不能靠提高模型温度或生成更多字段，而应补齐样例、规则、证据验证和离线回归，"
        "逐项提高可测的结果质量。",
    ]:
        add_body(doc, text)

    add_heading(doc, "8  结论与下一步计划")
    for text in [
        "本实践在 0.6 版本基础上完成了面向工程化使用的增量升级。系统没有改变原有 Rule First 主线，"
        "而是补齐了模型门控、企业资源范围、语义对象稳定性、对象级溯源、插件治理、配置中心、运行观测、"
        "基准测试和部署文档。升级后的系统能够将 PDF、DOCX 文档输出为带有页码、block、offset、"
        "规则和执行信息的标准 UOM Package，为上层应用提供可信的文档结构化输入。",
        "下一步计划包括：第一，针对扫描 PDF、表格和复杂版面增加可重复的标注样例，提升 OCR 后的标签与证据完整率；"
        "第二，补充 ResourcePlan 环境覆盖、异常回退和插件生命周期的测试，提高整体覆盖率；"
        "第三，逐步解决全库严格 mypy 的历史类型问题；第四，完善模型缓存校验和离线安装说明；"
        "第五，根据真实企业文档反馈扩展 YAML 标签和业务对象规则，但仍保持所有新增语义结果必须具备原文证据。",
        "总体而言，ZYRelay 的价值不在于自动替人做最终业务判断，而在于把文档中的关键信息稳定地提取出来，"
        "把每个结果准确地指回原文，并把处理过程中使用的配置和资源记录下来。这一能力是后续代码评审、"
        "知识服务、业务审核和智能应用可靠使用文档数据的基础。",
    ]:
        add_body(doc, text)

    add_heading(doc, "参考文献")
    references = [
        "[1] Python Software Foundation. Python Documentation[EB/OL]. https://docs.python.org/.",
        "[2] FastAPI. FastAPI Documentation[EB/OL]. https://fastapi.tiangolo.com/.",
        "[3] Pydantic. Pydantic Documentation[EB/OL]. https://docs.pydantic.dev/.",
        "[4] PyMuPDF. PyMuPDF Documentation[EB/OL]. https://pymupdf.readthedocs.io/.",
        "[5] python-docx. python-docx Documentation[EB/OL]. https://python-docx.readthedocs.io/.",
        "[6] YAML Language Development Team. YAML Specification[EB/OL]. https://yaml.org/.",
        "[7] PaddlePaddle. PaddleOCR Project Documentation[EB/OL]. https://github.com/PaddlePaddle/PaddleOCR.",
        "[8] ZYRelay DocIntelligence 项目文档：README、Architecture、ResourceManager、Benchmark、测试与基准记录[Z]. 2026.",
    ]
    for item in references:
        p = doc.add_paragraph()
        set_paragraph(p, after=2, first_line=0)
        r = p.add_run(item)
        set_font(r, 10.5)

    add_heading(doc, "附录 A  课题分工与实践体会")
    add_body(doc, "建议在最终提交前由团队补充真实分工。例如：成员 A 负责 PDF/DOCX 解析、统一 blocks 和标签规则；"
             "成员 B 负责资源插件、OCR 门控、模型回退与模型缓存；成员 C 负责语义对象、provenance、"
             "UOM 接口和 API；成员 D 负责测试、基准、Docker、CI 与文档。若为个人完成，可将本段改为个人在"
             "需求分析、编码、测试和报告撰写中的具体工作。")
    add_body(doc, "本次实践的直接体会是：面对非结构化数据，先建立“原文—位置—规则—结果”的最小闭环，"
             "再考虑模型增强和上层应用，能够降低系统复杂度。对于需要审计和复核的企业场景，"
             "能否定位原文通常比能否生成漂亮的总结更重要。")
    add_body(doc, "在提交和演示时，建议围绕一个完整样例说明系统能力：先上传合同或代码规范文档，展示解析后的页数和 blocks；"
             "再查询合同编号、金额或规范规则，展示其页码、block_id、原文片段和 offset；随后打开对应的 provenance，"
             "说明本次执行选择了哪些 Ground 配置与资源插件，是否发生 fallback；最后查看 UOM Package 中的 semantic_objects 和"
             "business_objects。这样的展示顺序与系统处理顺序一致，既能说明功能，也能避免把模型辅助结果误说成最终业务结论。"
             "对于出现 warning 的样例，应如实说明其原因，例如扫描件需要 OCR、模型资源不可用后使用启发式回退，"
             "并说明规则结果仍保留而没有被模型覆盖。")

    doc.core_properties.title = "ZYRelay DocIntelligence 实践报告"
    doc.core_properties.subject = "基于 0.6 版本的迭代优化实践"
    doc.core_properties.author = "待填写"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    # Chinese character count is deliberately checked for the assignment limit.
    from docx import Document as Reader
    text = "".join(p.text for p in Reader(path).paragraphs)
    print(path)
    print(f"characters={len(text)}")
