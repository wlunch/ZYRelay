from pathlib import Path

import fitz
import pytest
from docx import Document


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "contract.pdf"
    pdf = fitz.open()
    for page_text in [
        "采购合同\n\n合同编号：HT-2026-001\n甲方：北京甲方有限公司",
        "乙方：上海乙方有限公司\n\n合同金额：人民币100,000.00\n签订日期：2026年7月18日",
    ]:
        page = pdf.new_page()
        page.insert_textbox(
            fitz.Rect(72, 72, 520, 760),
            page_text,
            fontsize=12,
            fontname="china-s",
        )
    pdf.save(path)
    pdf.close()
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "contract.docx"
    document = Document()
    document.add_heading("采购合同", level=0)
    document.add_heading("合同信息", level=1)
    document.add_paragraph("合同编号：HT-2026-001")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "甲方"
    table.cell(0, 1).text = "北京甲方有限公司"
    table.cell(1, 0).text = "合同金额"
    table.cell(1, 1).text = "100000.00"
    document.add_paragraph("签订日期：2026年7月18日")
    document.save(path)
    return path


@pytest.fixture
def sample_convention_docx(tmp_path: Path) -> Path:
    path = tmp_path / "java-conventions.docx"
    document = Document()
    document.add_heading("后端团队 Java 开发规范", level=0)
    document.add_paragraph("适用于 Spring Boot 后端项目。")

    document.add_heading("1 命名规范", level=1)
    document.add_paragraph(
        "Java 类名必须使用大驼峰命名，例如 OrderService。"
        "禁止使用 order_service 作为类名。"
    )

    document.add_heading("2 方法长度", level=1)
    document.add_paragraph("单个方法不得超过 80 行。")

    document.add_heading("3 日志规范", level=1)
    document.add_paragraph(
        "禁止使用 System.out.println 输出业务日志，应使用统一日志框架。"
    )

    document.add_heading("4 安全规范", level=1)
    document.add_paragraph("不得在源代码中硬编码密码、Token 或数据库连接密钥。")

    document.add_heading("5 测试规范", level=1)
    document.add_paragraph("核心模块的单元测试覆盖率不得低于 80%。")

    document.add_heading("6 注释规范", level=1)
    document.add_paragraph("建议公共方法添加 Javadoc。")

    document.add_heading("7 评审要求", level=1)
    document.add_paragraph("所有变更必须经过 Code Review。", style="List Bullet")

    document.add_heading("8 工具规则表", level=1)
    table = document.add_table(rows=4, cols=3)
    values = [
        ("分类", "规范要求", "检查工具"),
        ("命名", "常量使用大写下划线", "Checkstyle"),
        ("格式", "单行不得超过120字符", "Checkstyle"),
        ("安全", "禁止提交明文密钥", "Semgrep"),
    ]
    for row, values_row in zip(table.rows, values, strict=True):
        for cell, value in zip(row.cells, values_row, strict=True):
            cell.text = value

    document.save(path)
    return path


@pytest.fixture
def sample_convention_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "java-conventions.pdf"
    pdf = fitz.open()
    for page_text in [
        "Java 团队代码规范\n\n1 命名规范\n\nJava 类名必须使用大驼峰命名。",
        "2 测试规范\n\n核心模块单元测试覆盖率不得低于 80%。",
    ]:
        page = pdf.new_page()
        page.insert_textbox(
            fitz.Rect(72, 72, 520, 760),
            page_text,
            fontsize=12,
            fontname="china-s",
        )
    pdf.save(path)
    pdf.close()
    return path
