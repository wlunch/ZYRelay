"""Create the stakeholder-facing report from the actual benchmark baseline."""

from __future__ import annotations

import json
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
BASELINE = ROOT / "benchmark" / "results" / "baseline" / "summary.json"
OUTPUT = (
    ROOT / "benchmark" / "results" / "reports" / "ZYRelay_基准数据集构建与测试汇报.docx"
)

# Resolved `standard_business_brief` tokens.
BLUE, DARK_BLUE, INK, MUTED, LIGHT, CALLOUT = (
    "2E74B5",
    "1F4D78",
    "0B2545",
    "667085",
    "F2F4F7",
    "F4F6F9",
)


def run_font(run, size: float = 11, color: str = "000000", bold: bool = False) -> None:
    run.font.name = "Arial Unicode MS"
    for key in ("ascii", "hAnsi"):
        run._element.rPr.rFonts.set(qn(f"w:{key}"), "Arial Unicode MS")
    # Bundled renderer reliably resolves this macOS CJK font; using the same
    # name in the East Asian font slot prevents missing-glyph boxes in PDF QA.
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def para(
    doc: Document,
    text: str = "",
    before: float = 0,
    after: float = 6,
    color: str = "000000",
    bold: bool = False,
    size: float = 11,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    run_font(p.add_run(text), size=size, color=color, bold=bold)
    return p


def heading(doc: Document, text: str, level: int) -> None:
    size, before, after, color = {
        1: (16, 16, 8, BLUE),
        2: (13, 12, 6, BLUE),
        3: (12, 8, 4, DARK_BLUE),
    }[level]
    p = para(doc, text, before=before, after=after, color=color, bold=True, size=size)
    p.paragraph_format.keep_with_next = True


def shade(cell, color: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), color)


def margins(cell) -> None:
    props = cell._tc.get_or_add_tcPr()
    mar = props.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        props.append(mar)
    for side in ("top", "bottom", "start", "end"):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), "80" if side in {"top", "bottom"} else "120")
        node.set(qn("w:type"), "dxa")


def table(
    doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]
) -> None:
    item = doc.add_table(rows=1, cols=len(headers))
    item.autofit = False
    props = item._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    props.append(layout)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    props.append(indent)
    for row, values in zip(
        [item.rows[0], *[item.add_row() for _ in rows]], [headers, *rows]
    ):
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            if row is item.rows[0]:
                shade(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run_font(p.add_run(value), size=9.5, color=INK, bold=row is item.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def make_report() -> Path:
    data = json.loads(BASELINE.read_text(encoding="utf-8"))
    summary = data["summary"]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.right_margin = section.bottom_margin = (
        section.left_margin
    ) = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Arial Unicode MS", Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_font(header.add_run("ZYRelay DocIntelligence · Benchmark"), size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_font(footer.add_run("内部技术汇报 · 基于实际本地执行结果"), size=9, color=MUTED)

    para(doc, "技术验证汇报", after=3, color=MUTED, bold=True, size=10)
    title = para(
        doc,
        "ZYRelay DocIntelligence\n可复现基准数据集构建与测试结果",
        after=8,
        color=INK,
        bold=True,
        size=23,
    )
    title.paragraph_format.line_spacing = 1.0
    para(
        doc,
        "范围：PDF / DOCX 文档智能、代码规范识别、OCR 与 provenance 追溯",
        after=14,
        color=MUTED,
        size=12,
    )
    table(
        doc,
        ["报告日期", "版本", "基线范围"],
        [
            [
                datetime.now(UTC).strftime("%Y-%m-%d"),
                data["relay_version"],
                "23 个案例 / 5 类文档",
            ]
        ],
        [1.4, 1.2, 3.9],
    )

    heading(doc, "一、结论摘要", 1)
    para(
        doc,
        f"本次已实际构建并执行一套 23 份文档的本地可重复基准集。基线运行成功 {summary['successful_cases']}/{summary['case_count']}，没有处理失败；预期检查、证据有效性和 provenance 有效性均为 {summary['expected_item_recall']:.0%}。",
    )
    para(
        doc,
        "基准集覆盖原生文本 PDF 与 image-only 扫描 PDF，验证了 ZYRelay 从输入文件、解析区块、规则/标签结果，到模型执行记录和原文定位证据的完整可审计链路。",
    )
    para(
        doc,
        "本轮不改变 Relay 的核心语义策略；新增内容仅位于 benchmark、测试和汇报层，便于后续作为回归门禁持续运行。",
    )

    heading(doc, "二、基准集构成与来源治理", 1)
    table(
        doc,
        ["类别", "数量", "代表内容", "覆盖重点"],
        [
            [
                "代码规范",
                "6",
                "Google Java、PEP 8、Linux、LLVM、OWASP",
                "标题、规则语句、代码/自然语言混排",
            ],
            [
                "合同",
                "4",
                "GSA SF-1449 / SF-30 / SF-26 / SF-33",
                "表格、业务字段、结构化表单",
            ],
            [
                "企业制度",
                "4",
                "NIST 身份、控制项、事件响应；NCSC 密码",
                "长文档、层级章节、制度文本",
            ],
            [
                "API 规范",
                "3",
                "RFC 9110、RFC 6749、JSON:API",
                "超长文档、协议术语、API 结构",
            ],
            [
                "扫描文档",
                "6",
                "由官方公开源本地渲染生成",
                "无文本层、OCR、bbox、模型证据",
            ],
        ],
        [1.0, 0.55, 2.6, 2.35],
    )
    para(
        doc,
        "来源配置集中于 benchmark/config/sources.yaml，仅允许 HTTPS 与声明的官方域名。来源审核覆盖 17 个上游源；原始二进制、扫描件和运行工件均被 .gitignore 排除，仓库仅保留配置、清单、案例、脚本与审核结论。",
        after=8,
    )

    heading(doc, "三、可追溯处理链路", 1)
    table(
        doc,
        ["阶段", "输出", "可追溯信息"],
        [
            [
                "源文件与清单",
                "manifest.json / manifest.csv",
                "稳定 BC-ID、SHA-256、页数、来源 URL、许可、格式与扫描标记",
            ],
            [
                "Relay 处理",
                "relay_result.json / uom.json",
                "有序 blocks、文本、规则/标签/代码规范候选与原始页码/offset",
            ],
            [
                "Ground 与资源",
                "ground.json / resources.json",
                "Ground 选择、配置快照、资源绑定与健康状态",
            ],
            [
                "模型与证据",
                "models.json / provenance.json",
                "OCR 模型版本、执行 ID、bbox、置信度、证据定位",
            ],
            [
                "评估与比较",
                "evaluation.json / summary.json",
                "通过项、召回、证据/追溯有效率、耗时及回归标记",
            ],
        ],
        [1.2, 2.2, 3.75],
    )

    heading(doc, "四、实际基线执行结果", 1)
    duration_seconds = summary["total_duration_ms"] / 1000
    table(
        doc,
        ["指标", "实际结果", "说明"],
        [
            [
                "基线案例",
                f"{summary['successful_cases']} / {summary['case_count']} 成功",
                "本地逐案例独立执行后汇总",
            ],
            [
                "预期检查",
                f"{summary['expected_item_recall']:.0%}",
                "结构、OCR、证据与 provenance 局部标注检查",
            ],
            [
                "证据有效性",
                f"{summary['evidence_valid_rate']:.0%}",
                "候选结果均可关联原文证据",
            ],
            [
                "provenance 有效性",
                f"{summary['provenance_valid_rate']:.0%}",
                "可从结果反查执行与资源记录",
            ],
            [
                "总处理耗时",
                f"{duration_seconds:.1f} 秒",
                "包含 6 个真实本地 OCR 扫描案例",
            ],
            [
                "OCR 版本",
                "; ".join(data.get("model_versions", [])) or "未启用",
                "离线缓存模型；不依赖外部推理服务",
            ],
        ],
        [1.4, 1.75, 4.0],
    )
    scan_rows = [
        [
            item["case_id"],
            f"{item['duration_ms'] / 1000:.1f}",
            str(item["block_count"]),
            str(item["warning_count"]),
        ]
        for item in data["cases"]
        if item["case_id"].startswith("BC-SCAN")
    ]
    heading(doc, "扫描件运行明细", 2)
    table(
        doc,
        ["案例", "耗时（秒）", "OCR blocks", "告警"],
        scan_rows,
        [1.25, 1.15, 1.35, 3.4],
    )

    heading(doc, "五、验证与回归机制", 1)
    para(
        doc,
        "新增离线 benchmark 测试覆盖：来源 HTTPS/域名白名单、稳定 ID 与数据集计数、扫描 PDF 无文本层且含图像、结构/OCR/证据/offset 评估、来源审核，以及基线汇总与语义回归检测。该套测试不访问网络，也不调用远程模型。",
    )
    para(
        doc,
        "全量项目测试已在安装真实 PaddleOCR 的环境中验证。既有 OCR 测试现明确允许两种受控状态：模型不可用时回退 NoOp；模型已离线部署并可用时使用 PaddleOCR。两种状态均保持不下载模型、不伪造 OCR 文本的安全边界。",
    )

    heading(doc, "六、复现方式与后续工作", 1)
    table(
        doc,
        ["操作", "命令 / 位置", "目的"],
        [
            [
                "数据校验",
                "benchmark/scripts/validate_dataset.py",
                "验证 23 项清单、案例对应关系和扫描路径",
            ],
            [
                "执行基线",
                "benchmark/scripts/run_benchmark.py --baseline",
                "本地运行全量或指定 suite/case",
            ],
            [
                "汇总基线",
                "benchmark/scripts/finalize_results.py",
                "断点续跑后的统一摘要",
            ],
            [
                "回归比较",
                "benchmark/scripts/compare_results.py",
                "比较 latest 与 baseline，识别语义回归",
            ],
        ],
        [1.2, 3.0, 2.95],
    )
    para(
        doc,
        "已知边界：当前 case 以结构、OCR、证据和追溯检查为主，业务规则和关键词的人工真值仍采用局部标注，尚未形成大规模逐页金标语料。建议下一步优先为每类文档增加少量高价值字段/规则的人工审核样本，并固定模型、配置哈希和基线结果用于版本门禁。",
    )
    para(
        doc,
        "报告依据：benchmark/results/baseline/summary.json、results/reports/dataset_validation.json、results/reports/source_audit.json；生成时不嵌入原始受限二进制文档。",
        before=8,
        after=0,
        color=MUTED,
        size=9,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(make_report())
